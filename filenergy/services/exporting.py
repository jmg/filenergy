"""Export a Conversation to PDF or DOCX.

PDF uses fpdf2 (pure Python, no Cairo); DOCX uses python-docx (already
required for indexing). Both functions return raw bytes so the view
layer can wrap them in a Flask Response.
"""
from __future__ import annotations

import io
import json
import logging
from typing import Iterable

log = logging.getLogger(__name__)


class ExportUnavailable(RuntimeError):
    pass


def _iter_messages(conversation) -> Iterable:
    return list(conversation.messages)


def _conversation_title(conversation) -> str:
    return conversation.title or f"Conversation {conversation.id}"


def _sources(msg) -> list[str]:
    if not msg.sources_json:
        return []
    try:
        items = json.loads(msg.sources_json)
    except Exception:
        return []
    return [str(item.get("name", "?")) for item in items if isinstance(item, dict)]


def to_markdown(conversation) -> str:
    """Same formatting as the /export.md endpoint, exposed as a service."""
    lines: list[str] = [f"# {_conversation_title(conversation)}", ""]
    if conversation.created_at:
        lines.append(f"_Created {conversation.created_at.strftime('%Y-%m-%d %H:%M')}_")
        lines.append("")
    for msg in _iter_messages(conversation):
        speaker = "**You**" if msg.role == "user" else "**Assistant**"
        lines.append(speaker)
        lines.append("")
        lines.append(msg.content or "")
        lines.append("")
        sources = _sources(msg)
        if sources and msg.role == "assistant":
            lines.append("Sources: " + ", ".join(sources))
            lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines)


def to_pdf(conversation) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ExportUnavailable("fpdf2 is not installed") from exc

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(18, 18, 18)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    # Explicit width so multi_cell knows where to wrap.
    text_w = pdf.w - pdf.l_margin - pdf.r_margin

    title = _conversation_title(conversation)
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(text_w, 8, _ascii(title))
    pdf.ln(2)

    if conversation.created_at:
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_text_color(120, 120, 120)
        pdf.multi_cell(
            text_w, 5,
            _ascii(f"Created {conversation.created_at.strftime('%Y-%m-%d %H:%M')}"),
        )
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

    for msg in _iter_messages(conversation):
        pdf.set_font("Helvetica", "B", 11)
        speaker = "You" if msg.role == "user" else "Assistant"
        pdf.multi_cell(text_w, 5, speaker)

        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(text_w, 5, _ascii(msg.content or " "))

        sources = _sources(msg)
        if sources and msg.role == "assistant":
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(120, 120, 120)
            pdf.multi_cell(text_w, 4, _ascii("Sources: " + ", ".join(sources)))
            pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

    out = pdf.output(dest="S")
    if isinstance(out, str):
        return out.encode("latin-1")
    if isinstance(out, bytearray):
        return bytes(out)
    return out


def to_docx(conversation) -> bytes:
    try:
        import docx
    except ImportError as exc:
        raise ExportUnavailable("python-docx is not installed") from exc

    document = docx.Document()
    document.add_heading(_conversation_title(conversation), level=1)
    if conversation.created_at:
        p = document.add_paragraph(
            f"Created {conversation.created_at.strftime('%Y-%m-%d %H:%M')}"
        )
        for run in p.runs:
            run.italic = True

    for msg in _iter_messages(conversation):
        speaker = "You" if msg.role == "user" else "Assistant"
        heading = document.add_paragraph()
        run = heading.add_run(speaker)
        run.bold = True

        document.add_paragraph(msg.content or "")

        sources = _sources(msg)
        if sources and msg.role == "assistant":
            sub = document.add_paragraph()
            sr = sub.add_run("Sources: " + ", ".join(sources))
            sr.italic = True

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _ascii(s: str) -> str:
    """fpdf2's default Helvetica is latin-1 only. Strip anything else."""
    return s.encode("latin-1", errors="replace").decode("latin-1")


# ---------------------------------------------------------------------------
# Workspace / user data export (GDPR portability)
# ---------------------------------------------------------------------------


def workspace_zip(workspace) -> bytes:
    """Bundle every workspace artifact into a ZIP and return its bytes.

    Layout:
        workspace.json            metadata (name, plan, members, policies)
        files/<id>_<name>         raw bytes from disk
        files.json                file metadata (size, indexed_at, summary, ...)
        conversations/<id>.md     per-conversation transcript
        conversations.json        thread metadata
        events.csv                audit log
        members.csv               members + roles
    """
    import csv
    import io
    import json
    import os
    import zipfile

    from filenergy.models import (
        Conversation, Event, File, WorkspaceMember,
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Top-level metadata
        zf.writestr("workspace.json", json.dumps({
            "id": workspace.id,
            "name": workspace.name,
            "slug": workspace.slug,
            "plan": workspace.plan,
            "require_2fa": bool(workspace.require_2fa),
            "owner_id": workspace.owner_id,
            "created_at": workspace.created_at.isoformat()
                if workspace.created_at else None,
        }, indent=2))

        # Members
        member_csv = io.StringIO()
        cw = csv.writer(member_csv)
        cw.writerow(["user_id", "email", "role", "joined_at"])
        for m in WorkspaceMember.query.filter_by(workspace_id=workspace.id):
            cw.writerow([
                m.user_id,
                m.user.email if m.user else "",
                m.role,
                m.created_at.isoformat() if m.created_at else "",
            ])
        zf.writestr("members.csv", member_csv.getvalue())

        # Files: bytes + metadata
        file_meta = []
        for f in File.query.filter_by(workspace_id=workspace.id):
            entry = {
                "id": f.id,
                "name": f.name,
                "size_bytes": f.size_bytes,
                "is_public": bool(f.is_public),
                "created_at": f.created_at.isoformat() if f.created_at else None,
                "indexed_at": f.indexed_at.isoformat() if f.indexed_at else None,
                "index_status": f.index_status,
                "summary": f.summary,
            }
            file_meta.append(entry)
            try:
                if f.path and os.path.isfile(f.path):
                    safe = "".join(
                        c if c.isalnum() or c in "._-" else "_"
                        for c in (f.name or "file")
                    )[:100]
                    zf.write(f.path, f"files/{f.id}_{safe}")
            except Exception:
                # Bytes missing on disk shouldn't kill the rest of the export.
                continue
        zf.writestr("files.json", json.dumps(file_meta, indent=2))

        # Conversations: per-thread Markdown + a manifest
        conv_meta = []
        for c in Conversation.query.filter_by(workspace_id=workspace.id):
            conv_meta.append({
                "id": c.id,
                "title": c.title,
                "user_id": c.user_id,
                "created_at": c.created_at.isoformat() if c.created_at else None,
                "messages": c.messages.count(),
            })
            zf.writestr(f"conversations/{c.id}.md", to_markdown(c))
        zf.writestr("conversations.json", json.dumps(conv_meta, indent=2))

        # Events
        event_csv = io.StringIO()
        ew = csv.writer(event_csv)
        ew.writerow(["created_at", "type", "user_id", "metadata_json"])
        for e in (
            Event.query.filter_by(workspace_id=workspace.id)
            .order_by(Event.id.asc())
        ):
            ew.writerow([
                e.created_at.isoformat() if e.created_at else "",
                e.type,
                e.user_id or "",
                e.metadata_json or "",
            ])
        zf.writestr("events.csv", event_csv.getvalue())

    return buf.getvalue()


def user_zip(user) -> bytes:
    """Per-user export — every workspace they own + their personal data."""
    import csv
    import io
    import json
    import zipfile

    from filenergy.models import (
        ApiKey, Conversation, Event, Workspace, WorkspaceMember,
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("user.json", json.dumps({
            "id": user.id,
            "email": user.email,
            "username": user.username,
            "totp_enabled": bool(user.totp_enabled),
            "weekly_digest": bool(user.weekly_digest),
            "created_at": user.created_at.isoformat() if user.created_at else None,
        }, indent=2))

        # Workspaces this user belongs to
        ws_meta = []
        for m in WorkspaceMember.query.filter_by(user_id=user.id):
            ws = m.workspace
            if ws is None:
                continue
            ws_meta.append({
                "id": ws.id,
                "name": ws.name,
                "slug": ws.slug,
                "role": m.role,
                "is_owner": ws.owner_id == user.id,
            })
            if ws.owner_id == user.id:
                # Bundle every owned workspace as a sub-archive.
                zf.writestr(
                    f"workspaces/{ws.slug}.zip",
                    workspace_zip(ws),
                )
        zf.writestr("workspaces.json", json.dumps(ws_meta, indent=2))

        # API keys (metadata only — plaintext was never stored)
        keys = ApiKey.query.filter_by(user_id=user.id).all()
        zf.writestr("api_keys.json", json.dumps([
            {
                "id": k.id,
                "name": k.name,
                "prefix": k.prefix,
                "scopes": k.scopes,
                "created_at": k.created_at.isoformat() if k.created_at else None,
                "revoked_at": k.revoked_at.isoformat() if k.revoked_at else None,
            }
            for k in keys
        ], indent=2))

        # Conversations the user authored across workspaces
        conv_csv = io.StringIO()
        cw = csv.writer(conv_csv)
        cw.writerow(["id", "workspace_id", "title", "messages", "created_at"])
        for c in Conversation.query.filter_by(user_id=user.id):
            cw.writerow([
                c.id, c.workspace_id, c.title or "",
                c.messages.count(),
                c.created_at.isoformat() if c.created_at else "",
            ])
            zf.writestr(f"conversations/{c.id}.md", to_markdown(c))
        zf.writestr("conversations.csv", conv_csv.getvalue())

        # Events for this user
        ev_csv = io.StringIO()
        ew = csv.writer(ev_csv)
        ew.writerow(["created_at", "type", "workspace_id", "metadata_json"])
        for e in Event.query.filter_by(user_id=user.id).order_by(Event.id.asc()):
            ew.writerow([
                e.created_at.isoformat() if e.created_at else "",
                e.type,
                e.workspace_id or "",
                e.metadata_json or "",
            ])
        zf.writestr("events.csv", ev_csv.getvalue())

    return buf.getvalue()
